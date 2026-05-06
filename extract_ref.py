import docx, os

base = r'd:\Project\aims\refrence\OpenClaw超级个体实操手册'
output = r'd:\Project\aims\ref_extract_super.md'
files = ['附录A 命令速查表.docx','附录B 常用Skills清单.docx','附录C 开箱即用的配置脚本模板.docx','附录F 安全防护指南.docx','附录G 国产Claw全景指南.docx']
sep = '=' * 80

with open(output, 'w', encoding='utf-8') as f:
    for fname in files:
        fpath = os.path.join(base, fname)
        f.write('\n' + sep + '\nFILE: ' + fname + '\n' + sep + '\n')
        try:
            doc = docx.Document(fpath)
            for para in doc.paragraphs:
                if para.text.strip():
                    f.write(para.text + '\n')
            for table in doc.tables:
                f.write('\n[TABLE]\n')
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    f.write(' | '.join(cells) + '\n')
                f.write('[/TABLE]\n')
        except Exception as e:
            f.write('[ERROR] ' + fname + ': ' + str(e) + '\n')
print('Done')
