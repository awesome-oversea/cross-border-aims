import docx
import os

base = r'd:\Project\aims\refrence\openclaw 权威指南配套素材与资源合集\书本配套可复制素材合集14个'
output = r'd:\Project\aims\ref_extract_book.md'
files = [
    '第2章OpenClaw 的技术架构.docx',
    '第4章本地安装步骤（Windows、Mac）.docx',
    '第12章  AI助手的记忆系统：它是如何记住你的.docx',
    '第13章  赋予AI人格与灵魂：OpenClaw从工具到伙伴转变的技巧.docx',
    '第14章 Skills技能系统：给AI添加超能力.docx',
    '第15章 多Agent团队协作系列教程（一）：搭建你的AI梦之队.docx',
    '第16章 多Agent团队协作系列教程（二）：使用子智能体实现协作过程清晰可见.docx',
    '第17章 多Agent团队协作系列教程（三）：两种协作方式如何选择.docx',
    '第18章 定时任务与自动化：让 AI在你睡觉的时候也能工作.docx',
]
sep = '=' * 80

with open(output, 'w', encoding='utf-8') as f:
    for fname in files:
        fpath = os.path.join(base, fname)
        f.write('\n' + sep + '\nFILE: ' + fname + '\n' + sep + '\n')
        try:
            doc = docx.Document(fpath)
            for para in doc.paragraphs:
                if para.text.strip():
                    f.write(para.text + '\n')
            for table in doc.tables:
                f.write('\n[TABLE]\n')
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    f.write(' | '.join(cells) + '\n')
                f.write('[/TABLE]\n')
        except Exception as e:
            f.write('[ERROR] ' + fname + ': ' + str(e) + '\n')
print('Done')
